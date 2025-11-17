import os
os.environ["CHROMA_TELEMETRY"] = "False"

import nltk
from nltk.corpus import stopwords
import docx
from pypdf import PdfReader
from dotenv import load_dotenv

from sentence_transformers import SentenceTransformer

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from chromadb import PersistentClient

from langchain_groq import ChatGroq
from langchain_core.documents import Document

load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

if not GROQ_API_KEY:
    raise ValueError("❌ Missing GROQ_API_KEY in .env")

nltk.download("stopwords")
nltk.download("punkt")

stop_words = set(stopwords.words("english"))


def preprocess_text(text):
    """Remove common English stopwords."""
    return " ".join([w for w in text.split() if w.lower() not in stop_words])

def load_pdf(path):
    reader = PdfReader(path)
    return "\n".join([(p.extract_text() or "") for p in reader.pages])


def load_docx(path):
    d = docx.Document(path)
    return "\n".join([p.text for p in d.paragraphs])


def load_txt(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def load_document(path):
    """Normalize and load file based on extension."""
    p = path.lower().strip().replace('"', "")  # remove quotes

    if p.endswith(".pdf"):
        return load_pdf(p)
    elif p.endswith(".docx"):
        return load_docx(p)
    elif p.endswith(".txt"):
        return load_txt(p)
    else:
        raise ValueError(f"Use PDF/DOCX/TXT only — Not supported: {p}")

# BUILD VECTOR DATABASE

def build_vector_db(path):

    raw = load_document(path)
    cleaned = preprocess_text(raw)

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=150
    )
    chunks = splitter.split_text(cleaned)
    docs = [Document(page_content=c) for c in chunks]

    embedder = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

    class Embedder:
        def embed_documents(self, texts):
            return embedder.encode(texts).tolist()

        def embed_query(self, text):
            return embedder.encode([text])[0].tolist()

    client = PersistentClient(path="db")

    vectordb = Chroma(
        client=client,
        collection_name="edugenie_collection",
        embedding_function=Embedder(),
    )

    vectordb.add_documents(docs)
    return vectordb


# BUILD RETRIEVER + LLM

def build_engines(vectordb):

    retriever = vectordb.as_retriever(search_kwargs={"k": 3})

    llm = ChatGroq(
        groq_api_key=os.getenv("GROQ_API_KEY"),
        model_name="llama-3.3-70b-versatile",
        temperature=0.1,
    )

    return retriever, llm
